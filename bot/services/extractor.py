"""Document text extraction service for Phase 3B.

Extracts plain text from PDF / DOCX / TXT / MD files for indexing as bookmarks.
Runs the (sync) parsing libraries inside ``asyncio.to_thread`` so the bot's
event loop stays responsive during slow PDFs.
"""
from __future__ import annotations

import asyncio
import logging
from dataclasses import dataclass
from pathlib import Path

logger = logging.getLogger(__name__)

# Hard cap on extracted characters. Long PDFs (e.g. books) can blow up the
# downstream AI pipeline / Telegram message limits otherwise.
MAX_CHARS = 50_000
_TRUNCATED_MARKER = "\n\n[обрезано]"


class ExtractError(Exception):
    """Raised when a document cannot be parsed."""


class EmptyDocumentError(ExtractError):
    """Document parsed successfully but contains no extractable text."""


class EncryptedPDFError(ExtractError):
    """PDF is password-protected — we don't try to decrypt."""


@dataclass(frozen=True)
class ExtractResult:
    text: str
    page_count: int | None
    truncated: bool


# ── Format dispatch ───────────────────────────────────────────


def _truncate(text: str) -> tuple[str, bool]:
    if len(text) <= MAX_CHARS:
        return text, False
    return text[:MAX_CHARS].rstrip() + _TRUNCATED_MARKER, True


def _extract_pdf_sync(path: Path) -> ExtractResult:
    from pypdf import PdfReader
    from pypdf.errors import PdfReadError

    try:
        reader = PdfReader(str(path))
    except PdfReadError as e:
        raise ExtractError(f"PDF повреждён: {e}") from e
    except Exception as e:
        raise ExtractError(f"Не удалось открыть PDF: {e}") from e

    if reader.is_encrypted:
        # Try empty password — many PDFs are "encrypted" with no password.
        try:
            if reader.decrypt("") == 0:
                raise EncryptedPDFError("PDF защищён паролем — не могу извлечь текст.")
        except EncryptedPDFError:
            raise
        except Exception as e:
            raise EncryptedPDFError(
                f"PDF защищён паролем — не могу извлечь текст: {e}"
            ) from e

    page_count = len(reader.pages)
    parts: list[str] = []
    total = 0
    for page in reader.pages:
        try:
            chunk = page.extract_text() or ""
        except Exception as e:
            logger.warning("PDF page extraction failed: %s", e)
            continue
        if not chunk:
            continue
        parts.append(chunk)
        total += len(chunk)
        if total >= MAX_CHARS:
            break

    text = "\n\n".join(parts).strip()
    if not text:
        raise EmptyDocumentError(
            "В PDF не удалось найти текст (возможно, это скан без OCR)."
        )
    truncated_text, truncated = _truncate(text)
    return ExtractResult(text=truncated_text, page_count=page_count, truncated=truncated)


def _extract_docx_sync(path: Path) -> ExtractResult:
    try:
        from docx import Document
    except ImportError as e:
        raise ExtractError("python-docx не установлен") from e

    try:
        doc = Document(str(path))
    except Exception as e:
        raise ExtractError(f"Не удалось открыть DOCX: {e}") from e

    parts: list[str] = []
    total = 0
    for para in doc.paragraphs:
        if not para.text:
            continue
        parts.append(para.text)
        total += len(para.text)
        if total >= MAX_CHARS:
            break

    # Also pull table cells — common in reports/CVs.
    if total < MAX_CHARS:
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if not cell.text:
                        continue
                    parts.append(cell.text)
                    total += len(cell.text)
                    if total >= MAX_CHARS:
                        break
                if total >= MAX_CHARS:
                    break
            if total >= MAX_CHARS:
                break

    text = "\n".join(parts).strip()
    if not text:
        raise EmptyDocumentError("DOCX пустой — не нашёл текста.")
    truncated_text, truncated = _truncate(text)
    return ExtractResult(text=truncated_text, page_count=None, truncated=truncated)


def _extract_plain_sync(path: Path) -> ExtractResult:
    # Best-effort decode: try utf-8, fall back to cp1251 (RU Windows), then latin-1.
    raw = path.read_bytes()
    if not raw:
        raise EmptyDocumentError("Файл пустой.")
    for encoding in ("utf-8", "utf-8-sig", "cp1251", "latin-1"):
        try:
            text = raw.decode(encoding)
            break
        except UnicodeDecodeError:
            continue
    else:
        raise ExtractError("Не удалось определить кодировку файла.")

    text = text.strip()
    if not text:
        raise EmptyDocumentError("Файл не содержит текста.")
    truncated_text, truncated = _truncate(text)
    return ExtractResult(text=truncated_text, page_count=None, truncated=truncated)


# ── Public API ────────────────────────────────────────────────


def detect_format(mime_type: str | None, filename: str | None) -> str | None:
    """Return one of {'pdf', 'docx', 'plain'} or None if unsupported."""
    mime = (mime_type or "").lower()
    name = (filename or "").lower()

    if mime == "application/pdf" or name.endswith(".pdf"):
        return "pdf"
    if (
        mime
        == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        or name.endswith(".docx")
    ):
        return "docx"
    if mime in ("text/plain", "text/markdown") or name.endswith((".txt", ".md")):
        return "plain"
    return None


async def extract_text(path: Path, fmt: str) -> ExtractResult:
    """Extract text from ``path``. ``fmt`` is one of {'pdf', 'docx', 'plain'}."""
    if not path.exists():
        raise ExtractError(f"Файл не найден: {path.name}")
    if path.stat().st_size == 0:
        raise EmptyDocumentError("Файл пустой.")

    if fmt == "pdf":
        return await asyncio.to_thread(_extract_pdf_sync, path)
    if fmt == "docx":
        return await asyncio.to_thread(_extract_docx_sync, path)
    if fmt == "plain":
        return await asyncio.to_thread(_extract_plain_sync, path)
    raise ExtractError(f"Неподдерживаемый формат: {fmt}")
